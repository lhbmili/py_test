from docx import Document
import os

OLDPATH = "D:\\GitHub program\\py_test\\未完成"
PATH = "D:\\GitHub program\\py_test\\new"

DICT = {
    "（点击展示条款内容）":"",
    "（点击导航）":"",
    "点击展示":"详见",
    "(点击链接展示流程图)":"（详见《税务行政职权运行流程图》）",
    # "税务行政职权运行流程图":"（详见《税务行政职权运行流程图》）",
    "（点击链接至办税指南该具体事项适用条件）":""
    }

def main():

    for fileName in os.listdir(OLDPATH):
        oldFile = OLDPATH + "\\" + fileName
        newFile = PATH + "\\" + fileName
        document = Document(oldFile)
        document = check(document,fileName)
        document.save(newFile)
        # 判断是否为docx
        # if oldFile.split(".")[1] == 'docx':
        #     document = Document(oldFile)
        #     document = check(document)
        #     document.save(newFile)   

def check(document,fileName):
    # tables
    for table in document.tables:
        for row in range(len(table.rows)):
            for col in range(len(table.columns)):
                # for key, value in DICT.items():
                #     if key in table.cell(row ,col).text:
                for para in table.cell(row ,col).paragraphs:
                    for i in range(len(para.runs)):
                        for key1, value1 in DICT.items():
                            if key1 in para.runs[i].text:
                                print("%s中有改动！"%fileName+key1+"->"+value1)
                                para.runs[i].text = para.runs[i].text.replace(key1, value1)

    # paragraphs
    for para in document.paragraphs:
        for i in range(len(para.runs)):
            for key, value in DICT.items():
                if key in para.runs[i].text:
                    print("filename:%s"%fileName+key+"->"+value)
                    para.runs[i].text = para.runs[i].text.replace(key, value)

    return document 
    
if __name__ == '__main__':
	main()