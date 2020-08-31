import copy
import docx
import os
import sys


def getFilePath(dic):
    """
    取得文件路径
    """
    file_paths = []
    s = os.path.join(sys.path[0], dic)
    list1 = os.listdir(s)
    for i in list1:
        file_paths.append("".join((s,"\\",i)))
    return (file_paths,i)

def change_key(doc, new_doc, dic):
  if dic and doc and new_doc:
    doc_path = os.path.join(FILE_PATH, doc)
    new_doc_path = os.path.join(NEW_FILE_PATH, new_doc)
    doc_file = docx.Document(doc_path)
    # 替换段落内容
    for paragraph in doc_file.paragraphs:
        # 深度复制段落内容，包括样式。如果不深度复制，样式会丢失
        list_runs = copy.deepcopy(paragraph.runs)
        paragraph.clear()
        #文字替换
        for run in list_runs:
            for name in dic:
                print (name)                
                if name in run.text:
                    value = dic[name]
                    run.text = run.text.replace(name, str(value))
            # 段落样式的复制
            paragraph.add_run(run.text, run.style)

    # 替换表格内容
    for table in doc_file.tables:
        # 深度复制表格内容，包括样式, 如果不深度复制，样式会丢失
        table_style = copy.deepcopy(table.style)
        for row in table.rows:
            for cell in row.cells:
                for name in dic:
                    if name in cell.text:
                        value = dic[name]
                        cell.text = cell.text.replace(name, str(value))
        table.style=table_style
    doc_file.save(new_doc_path)
    print(new_doc_path)

if __name__ == "__main__":
    replace_text = {"（点击展示条款内容）":"","（点击导航）":"","点击展示":"详见"}
    files = getFilePath("2链接信息表 （第二税务分局）")
    for path,name in files:
        print(path,name)
        # doc = Document(path)
        # replaceStr(doc,replace_text)
        # doc.save("%s" % name)